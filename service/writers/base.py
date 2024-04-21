from abc import ABC, abstractmethod

from docx.shared import RGBColor

from service.banwords import is_banword
from service.phones import strong_phones
from service.links.base import Linker

RED = RGBColor(255, 0, 0)


class BasePageWriter(ABC):
    def __init__(
        self,
        document,
        data: dict,
        banwords: list[str] | None = None,
        phone: str | None = None,
        linker: Linker | None = None,
    ):
        self.banwords = banwords or []
        self.phone = phone
        self.document = document
        self.data = data
        self.linker = linker

    def write_title(self, title: str) -> None:
        self.document.add_heading(title, level=0)

    def write_heading(self, heading: str, level: int) -> None:
        text = self._wrap_tag(f'h{level}', heading)
        self.document.add_heading(text, level=level)

    def insert_links(self, text: str) -> str:
        if not self.linker:
            return text
        return self.linker.render(text)

    def write_paragraph(
        self, 
        text: str, 
        make_phone_bold: bool = False, 
        insert_links: bool = False,
    ) -> None:
        if make_phone_bold:
            text = self._make_phone_bold(text)

        if insert_links:
            text = self.insert_links(text)

        paragraph = self.document.add_paragraph('<p>')
        words = text.split()
        for word in words:
            word = word + ' '
            if word.startswith('<p>') and not make_phone_bold:
                word = '\n' + word
            if is_banword(word.lower(), self.banwords):
                run = paragraph.add_run(word)
                run.font.color.rgb = RED
            elif '<strong>' in word and '</strong>' in word:
                run = paragraph.add_run(word)
                run.bold = True
            else:
                paragraph.add_run(word)
        self.document.add_paragraph('</p>')

    def _make_phone_bold(self, text: str) -> str:
        if not self.phone:
            return text
        return strong_phones(text, self.phone)

    def _wrap_tag(self, tag: str, text: str) -> str:
        return f'<{tag}>{text}</{tag}>'

    def add_page_break(self) -> None:
        self.document.add_page_break()

    @abstractmethod
    def write(self) -> None:
        raise NotImplementedError('You must implement write method')
